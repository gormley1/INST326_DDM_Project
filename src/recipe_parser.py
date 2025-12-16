"""
Recipe Parser Inheritance Hierarchy
INST326 - Fall 2025
Team Member: Darrell Cox - Documentation Lead

Implements an inheritance hierarchy for parsing recipes from different file formats.
Supports .txt, .pdf, .docx formats with unified ingredient extraction and structure.
"""

import os
import re
from abc import ABC, abstractmethod
from typing import Dict, List
import PyPDF2  # Must be installed
# python-docx imported lazily inside DOCX parser


# ======================================================================
#                           BASE CLASS
# ======================================================================

class RecipeParser(ABC):
    """Abstract base class for all recipe parsers."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.recipe_data = {}

    @abstractmethod
    def parse(self) -> Dict:
        """Parse the file into structured recipe data."""
        pass

    @abstractmethod
    def validate_format(self) -> bool:
        """Check if the file exists and is the expected format."""
        pass

    # ---------- Shared Ingredient Parsing Utilities ----------

    def extract_ingredients_section(self, text: str) -> List[str]:
        """
        Extract ingredient lines from recipe text, including table formats.
        
        Enhanced for Project 4 to handle:
        - Traditional bullet list format
        - Table format with "Ingredient | Weight | Measure" columns
        - Institutional/commerical recipe formats

        Args:
            text (str): Raw recipe text

        Returns:
            List[str]: Cleaned ingredient lines
        """
        lines = text.split("\n")
        ingredients = []
        in_section = False

# https://www.w3schools.com/python/gloss_python_regex_metacharacters.asp 
        # Patterns for detecting ingredient section start
        ingredients_headers = [r"^ingredient?:?$", 
                               r"^ingredient\s+weight\s+measure", 
                               r"^ingredient$"]
        # Patterns for detecting section end
        end_patterns = [r"^directions?|instructions?|steps?|method?|):?$",
                        r"^\s*method\s*$",
                        r"^\s*\d+\s*$"] # if the steps are numbered
        # Words that indicate we're NOT in ingredients section anymore
        stop_words = ['method', 'directions', 'instructions', 'steps', 'calories', 'yield', 'portion', 'nutrition']

        for i, line in enumerate(lines):
            clean = line.strip()
            #if entering ingredients section
            for pattern in ingredients_headers:
                if re.match(pattern, clean, re.IGNORECASE):
                    in_section = True
                    continue
            # if leaving ingredients section
            if in_section: 
                for pattern in end_patterns:
                    if re.match(pattern, clean, re.IGNORECASE):
                        in_section = False
                        break
                if any(word in clean.lower() for word in stop_words):
                    in_section = False
                    continue
            # extract ingredient if in section
            if in_section and clean:
                skip_lines = ['ingredient', 'weight', 'measure', 'issue', 'quantity', 'unit', 'amount']
                if clean.lower() in skip_lines:
                    continue
            clean = re.sub(r"[\-~+•*◦▪▫→]\s*|>>\s*|-->\s*|->\s*", "", clean) 

            # this block only extracts ingredient name part for table format; 
            # leaving it commented out bc I worry it'll break the ingredient measurement + summation

            # if the line has measureents, after ingredient name, extract just the ingredient part
            """ if any(unit in clean for unit in [' lbs', ' oz', ' gal', ' cup', ' tbsp', ' tsp']):
                parts = re.split(r'\s+\d+[\-\d/]*\s+(lbs?|oz|gal|cup|tbsp|tsp)', clean)
                if parts and len(parts[0]) > 3:
                    clean = parts[0].strip()
            """
            if len(clean) > 3:
                ingredients.append(clean)

        return ingredients

    def clean_ingredient_text(self, text: str) -> str:
        """Remove bullet characters & trim whitespace."""
        text = re.sub(r"^[\-•*◦▪▫→]\s*", "", text)
        return " ".join(text.split()).strip()

    # ---------- Convenience Accessors ----------

    def get_recipe_name(self) -> str:
        return self.recipe_data.get("name", "Untitled Recipe")

    def get_ingredients(self) -> List[str]:
        return self.recipe_data.get("ingredients", [])


# ======================================================================
#                         TXT PARSER
# ======================================================================

class TXTRecipeParser(RecipeParser):

    def validate_format(self) -> bool:
        return self.filepath.endswith(".txt") and os.path.isfile(self.filepath)

    def parse(self) -> Dict:
        if not self.validate_format():
            raise ValueError(f"Invalid TXT file: {self.filepath}")

        with open(self.filepath, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.split("\n")
        name = lines[0].strip() if lines else "Untitled Recipe"

        ingredients = [self.clean_ingredient_text(i)
                       for i in self.extract_ingredients_section(content)]

        # Directions extraction
        directions = []
        in_directions = False
        for line in lines:
            line_clean = line.strip()
            if re.match(r"^(directions?|instructions?|steps?):?$", line_clean, re.IGNORECASE):
                in_directions = True
                continue
            if in_directions and line_clean:
                directions.append(line_clean)

        self.recipe_data = {
            "name": name,
            "ingredients": ingredients,
            "directions": directions,
            "source_file": self.filepath,
            "format": "txt",
        }

        return self.recipe_data


# ======================================================================
#                         PDF PARSER
# ======================================================================

class PDFRecipeParser(RecipeParser):

    def validate_format(self) -> bool:
        if not self.filepath.endswith(".pdf") or not os.path.isfile(self.filepath):
            return False
        try:
            with open(self.filepath, "rb") as f:
                PyPDF2.PdfReader(f)
            return True
        except Exception:
            return False

    def parse(self) -> Dict:
        """Parses a PDF file into recipe data w/ broader functionality for ingredient extraction"""
        if not self.validate_format():
            raise ValueError(f"Invalid PDF file: {self.filepath}")

        full_text = ""

        with open(self.filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    full_text += txt + "\n"

        lines = [l.strip() for l in full_text.split("\n") if l.strip()]

        # ----- added during debugging: looking for recipe title in first few lines, extracting name -----
        name = "Untitled Recipe"
        for line in lines[:5]:
            # this might cause some bugs but users can rename it anyways if it's a problem
            if any(word in line.lower() for word in ['dairy', 'meat', 'poultry', 'no.', 'yield', 'portion']):
                continue
            if line.isupper() or line.istitle():
                if 10 < len(line) < 100:
                    name = line 
                    break
        
        if name == "Untitled Recipe" and lines:
            name = lines[0]

        ingredients = self.extract_ingredients_section(full_text)



        # ------------------------------------------------------------------------------------------------

        # Directions
        directions = []
        in_directions = False
        for line in lines:
            if re.match(r"^(directions?|instructions?|steps?):?$", line, re.IGNORECASE):
                in_directions = True
                continue
            if in_directions:
                if any(word in line.lower() for word in ['calories', 'nutriotion', 'yield']):
                    break
                if line and not line.isdigit():
                    directions.append(line)

        self.recipe_data = {
            "name": name,
            "ingredients": ingredients,
            "directions": directions,
            "source_file": self.filepath,
            "format": "pdf",
        }

        return self.recipe_data


# ======================================================================
#                         DOCX PARSER
# ======================================================================

class DOCXRecipeParser(RecipeParser):

    def validate_format(self) -> bool:
        if not self.filepath.endswith(".docx") or not os.path.isfile(self.filepath):
            return False
        try:
            from docx import Document
            Document(self.filepath)
            return True
        except Exception:
            return False

    def parse(self) -> Dict:
        if not self.validate_format():
            raise ValueError(f"Invalid DOCX file: {self.filepath}")

        from docx import Document
        doc = Document(self.filepath)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        name = doc.paragraphs[0].text.strip() if doc.paragraphs else "Untitled Recipe"

        ingredients = [self.clean_ingredient_text(i)
                       for i in self.extract_ingredients_section(full_text)]

        # Directions
        lines = full_text.split("\n")
        directions = []
        in_directions = False
        for line in lines:
            clean = line.strip()
            if re.match(r"^(directions?|instructions?|steps?):?$", clean, re.IGNORECASE):
                in_directions = True
                continue
            if in_directions and clean:
                directions.append(clean)

        self.recipe_data = {
            "name": name,
            "ingredients": ingredients,
            "directions": directions,
            "source_file": self.filepath,
            "format": "docx",
        }

        return self.recipe_data


# ======================================================================
#                        DEMO / MAIN USAGE
# ======================================================================

def demo_recipe_parsers():
    """Showcase all parser types working with inheritance & polymorphism."""

    print("=== Demo: TXT Parser ===")
    txt = TXTRecipeParser("data/sample_recipes/pasta_marinara.txt")
    if txt.validate_format():
        r = txt.parse()
        print("Parsed:", r["name"])

    print("\n=== Demo: PDF Parser ===")
    pdf = PDFRecipeParser("data/sample_recipes/chicken_stir_fry.pdf")
    if pdf.validate_format():
        r = pdf.parse()
        print("Parsed:", r["name"])

    print("\n=== Polymorphism Demo ===")
    for parser in [txt, pdf]:
        print(f"{parser.get_recipe_name()} ({parser.recipe_data.get('format')})")


def usage_in_main_project():
    """Example how the main DDM grocery system could consume the parsers."""
    print("\n=== Integration Example ===")

    recipe_files = [
        "data/sample_recipes/pasta_marinara.txt",
        "data/sample_recipes/chicken_stir_fry.pdf",
    ]

    recipes = []

    for filepath in recipe_files:
        if filepath.endswith(".txt"):
            parser = TXTRecipeParser(filepath)
        elif filepath.endswith(".pdf"):
            parser = PDFRecipeParser(filepath)
        elif filepath.endswith(".docx"):
            parser = DOCXRecipeParser(filepath)
        else:
            print("Unsupported format:", filepath)
            continue

        if parser.validate_format():
            recipes.append(parser.parse())
            print("✓ Parsed:", parser.get_recipe_name())
        else:
            print("✗ Failed:", filepath)

    print("\nTotal Recipes:", len(recipes))


if __name__ == "__main__":
    demo_recipe_parsers()
    usage_in_main_project()
